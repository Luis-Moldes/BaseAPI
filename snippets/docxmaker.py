from docx.shared import Cm
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import datetime as dt
import copy

#import source.rexywind.mathfunctions as math_functions


def get_section_max_from_dictionary(dictionary):
    maxi = None
    for section in dictionary:

        if maxi is None:
            maxi = dictionary[section]['number']
        else:
            maxi = max(maxi, dictionary[section]['number'])
    return maxi


def get_graph_number_max_from_dictionary(dictionary):
    maxi = None
    for section in dictionary:
        for graph_name in dictionary[section]['graph']:
            if maxi == None:
                maxi = dictionary[section]['graph'][graph_name]['graph_number']
            else:
                maxi = max(maxi, dictionary[section]['graph'][graph_name]['graph_number'])
    return maxi


def add_focus_to_template(dictionary, race, leg, section_number=None,
                          preset_path=None):  # legs are tuple, not user friendly do something
    if type(race) == type(1):
        race = race,
    if section_number == None:
        section_max = get_section_max_from_dictionary(dictionary)
        section_key = 'section' + str(section_max + 1)
        dictionary[section_key]['graph'] = defaultdict(dict)
    else:
        section_key = 'section' + str(section_number)
        dictionary[section_key]['graph'] = defaultdict(dict)
    graph_number = get_graph_number_max_from_dictionary(dictionary)
    for i_race in race:
        for i_leg in leg:
            graph_number = graph_number + 1
            chart_name = 'Chart_Focus_Race' + str(i_race) + 'Leg' + str(i_leg)
            dictionary[section_key]['graph'][chart_name]['title'] = None  # 'Race ' + str(i_race) + ' Leg ' + str(i_leg)
            dictionary[section_key]['graph'][chart_name]['graph_number'] = graph_number
            dictionary[section_key]['graph'][chart_name]['type'] = 'chart_focus'
            dictionary[section_key]['graph'][chart_name]['preset'] = preset_path
            dictionary[section_key]['graph'][chart_name]['path'] = 'Focus.png'
    return dictionary


def add_recursive_element_to_template(dictionary, race_number, race_graph_name, race_type, race_preset,
                                      leg_number, leg_graph_name, leg_type, leg_preset):
    if not isinstance(race_number, tuple):
        race_number = race_number,
    if not isinstance(race_graph_name, tuple):
        race_graph_name = race_graph_name,
    if not isinstance(race_type, tuple):
        race_type = race_type,
    if not isinstance(race_preset, tuple):
        race_preset = race_preset,
    if not isinstance(leg_number, tuple):
        leg_number = leg_number,
    if not isinstance(leg_graph_name, tuple):
        leg_graph_name = leg_graph_name,
    if not isinstance(leg_type, tuple):
        leg_type = leg_type,
    if not isinstance(leg_preset, tuple):
        leg_preset = leg_preset,
    out = False
    if len(race_number) != len(leg_number):
        print("Wrong race_number or leg_number input, size not matching")
        out = True
    if len(race_graph_name) != len(race_preset):
        print("Size not matching in the template file between race_graph_name and race_preset")
        out = True
    if len(race_graph_name) != len(race_type):
        print("Size not matching in the template file between race_graph_name and race_type")
        out = True
    if len(leg_graph_name) != len(leg_preset):
        print("Size not matching in the template file between leg_graph_name and leg_preset")
        out = True
    if len(leg_graph_name) != len(leg_type):
        print("Size not matching in the template file between leg_graph_name and leg_type")
        out = True
    if out:
        exit()

    section_number = get_section_max_from_dictionary(dictionary)
    graph_number = get_graph_number_max_from_dictionary(dictionary)
    if section_number is None:
        section_number = 0
    if graph_number is None:
        graph_number = 0
    race_index = 0
   # deal race leg graphs
    for i_race in race_number:
        section_key = 'section' + str(section_number + 1)
        dictionary[section_key]['graph'] = defaultdict(dict)
        ii = 0
        if (race_graph_name is not None) & (leg_number[race_index] > 1):
            for element in race_type:
                graph_number = graph_number + 1
                graph_name = race_graph_name[ii] + '_Race' + str(i_race)
                dictionary[section_key]['graph'][graph_name]['graph_number'] = graph_number
                dictionary[section_key]['graph'][graph_name]['type'] = element
                dictionary[section_key]['graph'][graph_name]['preset'] = race_preset[ii]
                dictionary[section_key]['graph'][graph_name]['path'] = None
                if ii == 0:
                    dictionary[section_key]['graph'][graph_name]['title'] = 'Race ' + str(i_race)
                else:
                    dictionary[section_key]['graph'][graph_name]['title'] = None
                dictionary[section_key]['graph'][graph_name]['filter'] = 'R' + str(i_race)
                ii = ii + 1

        for i_leg in range(1, leg_number[race_index] + 1):
            for jj in range(0, len(leg_graph_name)):
                graph_number = graph_number + 1
                graph_name = leg_graph_name[jj] + '_Race' + str(i_race) + 'Leg' + str(i_leg)
                dictionary[section_key]['graph'][graph_name]['graph_number'] = graph_number
                dictionary[section_key]['graph'][graph_name]['type'] = leg_type[jj]
                dictionary[section_key]['graph'][graph_name]['preset'] = leg_preset[jj]
                dictionary[section_key]['graph'][graph_name]['title'] = None  # 'Race'+str(i_race)+' Leg'+str(i_leg)
                dictionary[section_key]['graph'][graph_name]['filter'] = 'R' + str(i_race) + 'L' + str(i_leg)
                dictionary[section_key]['graph'][graph_name]['path'] = None

        race_index = race_index + 1
        section_number = section_number + 1
    return dictionary


def clean_docx(path):
    document = Document(path)
    for paragraph in document.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # changing the page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)  # Cm(1.6)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.6)
    return document


def add_header(document, title):
    section = document.sections[0]
    header = section.header
    paragraph = header.add_paragraph()
    paragraph.text = title
    header.add_paragraph('                                ')
    return document

def create_word_document(template_dictionary, template_path, boatdata):
    print("Start writing word document")
    document = Document(template_path)
    # document = clean_docx(template_path)
    counter = 1
    boolean = False
    for section in template_dictionary:
        print("Section " + str(counter))
        section_graph_number = 0
        for graph in template_dictionary[section]['graph']:
            section_graph_number = section_graph_number + 1
            print(graph)
            if template_dictionary[section]['graph'][graph]['title'] is not None:
                document.add_heading(template_dictionary[section]['graph'][graph]['title'], 1)
            if template_dictionary[section]['graph'][graph]['path'] is not None:
                document.add_picture(template_dictionary[section]['graph'][graph]['path'], width=Cm(19.7))
            if ('section' + str(counter + 1)) in template_dictionary:
                next_section_graph_number = len(template_dictionary['section' + str(counter + 1)]['graph'])
                boolean = next_section_graph_number > 1
            else:
                boolean = False
        if (counter < len(template_dictionary)) & ((section_graph_number > 1) | boolean):
            document.add_page_break()
            counter = counter + 1
    title = 'Wind - ' + boatdata.iloc[0]['BoatName'] + " - " + boatdata.iloc[0]['Date']
    add_header(document, title)
    document.save('Output/' + title + ".docx")

# Solution to add header and footer

# msword = win32.gencache.EnsureDispatch('Word.Application')
# tempfile = r"C:\Users\jlepi\Documents\python_project\windreport\test.docx"  #full path
# doc = msword.Documents.Open(tempfile)
# doc.Sections(1).Headers(1).Range.Text = r'Text to be included'
# doc.Sections(1).Headers(1).Range..AddPicture(r"C:\Users\jlepi\Documents\python_project\windreport\KND_logo.png")
# doc.Sections(1).Footers(1).PageNumbers.Add()
# tempfile = r"C:\Users\jlepi\Documents\python_project\windreport\test.docx"
# #doc.SaveAs(tempfile, FileFormat = 0)
#
# #document = Document(tempfile)
#


def add_reach_dn_focus_preset(data, event, template, dn_path, reach_path):
    for section in template:
        for graph in template[section]['graph']:
            if template[section]['graph'][graph]['type'] == 'chart_focus':
                filtered_data, events_to_plot = math_functions. \
                    data_filter_from_template_info(template, section, graph, data, event)
                sailing_mode = filtered_data.SailingMode.value_counts().index[0]
                if sailing_mode == 'D':
                    template[section]['graph'][graph]['preset'] = dn_path
                if sailing_mode == 'O':
                    template[section]['graph'][graph]['preset'] = reach_path
    return template


def quickhacktogetadocx(graphlist, templatepath, boatname, date, tables):
    document = Document(templatepath)
    i=0
    for graph in graphlist:
         if i==0:
             if len(tables)!=1:
                 document.add_paragraph('        ')
                 document.add_picture(graph, width=Cm(12))
                 last_paragraph = document.paragraphs[-1]
                 last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
             else:
                 document.add_paragraph('        ')
                 document.add_picture(graph, width=Cm(15))
                 last_paragraph = document.paragraphs[-1]
                 last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

             if len(tables) != 1:
                 for tab in tables:
                     document.add_picture(tab, width=Cm(14))
                     last_paragraph = document.paragraphs[-1]
                     last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
             else:
                 document.add_picture(tables[0], width=Cm(16))
                 last_paragraph = document.paragraphs[-1]
                 last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

         elif i == 1:
             document.add_picture(graph, width=Cm(12))
             last_paragraph = document.paragraphs[-1]
             last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

         elif (i>1 and i<8):
             document.add_picture(graph, width=Cm(16))
             last_paragraph = document.paragraphs[-1]
             last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
             document.add_paragraph('        ')
         else:
             document.add_picture(graph, width=Cm(11.5))
             last_paragraph = document.paragraphs[-1]
             last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
             document.add_paragraph('        ')
             if i%2 == 0:
                 document.add_paragraph('        ')
         i+=1
    title = "GOLD Report - " + boatname + " - " + date
    add_header(document, title)
    document.save('Outputs/' + title + ".docx")
    print("Report saved")
    return 'Outputs/' + title + ".docx"


def quickhacktogetadocx_rt(graphlist, templatepath, boatname, date):
    document = clean_docx(templatepath)
    i=0
    for graph in graphlist:
        if i==0:
            document.add_paragraph('        ')
            document.add_picture(graph, width=Cm(15))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif i==1:
            document.add_picture(graph, width=Cm(14))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = last_paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)

        else:
            document.add_picture(graph, width=Cm(9))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph('        ')

        i=i+1
    title = "GOLD Analysis Parameter Finder - " + boatname + " - " + date
    add_header(document, title)
    document.save('Outputs/Remember/' + title + ".docx")
    print("report saved")

    return 'Outputs/Remember/' + title + ".docx"


def quickhacktogetadocx_API(num, root, sq, templatepath, boatname, date):
    document = clean_docx(templatepath)
    document.add_paragraph('Your number is ' + str(num))
    document.add_paragraph('Your square root is ' + str(root))
    document.add_paragraph('Your square is ' + str(sq))
    title = "GOLD's Incredible Calculator - " + str(boatname) + " - " + date
    add_header(document, title)
    document.save('media/Files/' + title + ".docx")

    return 'Files/' + title + ".docx"


def addlineup(dictionary, event):
    dico = copy.deepcopy(dictionary)
    for section in dictionary:
        for graph in dictionary[section]['graph']:
            if dictionary[section]['graph'][graph]['type'] == 'line_up':
                ilineup = 1
                for lineup in event.itertuples(index=False):
                    start_time = lineup.DateTime
                    stop_time = lineup.DateTime + dt.timedelta(0, int(lineup.attribute))
                    graphname = 'Line_up' + str(ilineup)
                    dico[section]['graph'][graphname]['type'] = 'chart_focus_multi'
                    n = int(dictionary[section]['graph'][graph]['graph_number'])
                    dico[section]['graph'][graphname]['graph_number'] = ilineup + n - 1
                    if ilineup == 1:
                        dico[section]['graph'][graphname]['title'] = 'Line Up'
                    else:
                        dico[section]['graph'][graphname]['title'] = None
                    dico[section]['graph'][graphname]['path'] = None
                    preset = dictionary[section]['graph'][graph]['preset']
                    dico[section]['graph'][graphname]['preset'] = preset
                    dico[section]['graph'][graphname]['filter'] = 'time'
                    dico[section]['graph'][graphname]['startfilter'] = start_time
                    dico[section]['graph'][graphname]['stopfilter'] = stop_time
                    ilineup = ilineup + 1
    return dico


